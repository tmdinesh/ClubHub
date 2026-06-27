from __future__ import annotations

import io
from datetime import date, datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colour palette matching the HTML preview ──────────────────────────────────
_NAVY      = RGBColor(0x1A, 0x1A, 0x2E)   # header bg / title text
_INDIGO    = RGBColor(0x31, 0x2E, 0x81)   # section headings
_SLATE     = RGBColor(0x47, 0x55, 0x69)   # kv label text
_BODY      = RGBColor(0x1E, 0x29, 0x3B)   # body text
_MUTED     = RGBColor(0x64, 0x74, 0x8B)   # sub-labels
_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_FG  = RGBColor(0x16, 0xA3, 0x4A)
_BLUE_FG   = RGBColor(0x1D, 0x4E, 0xD8)

_HEX_NAVY    = "1A1A2E"
_HEX_HEADER_BG = "1E1B4B"   # deep indigo for stat-box headers
_HEX_KV_LABEL  = "F1F5F9"   # light blue-grey for kv label cells
_HEX_STAT_BOX  = "F8FAFC"   # stat box background
_HEX_EVENT_BG  = "F8FAFC"   # event block background
_HEX_GREEN_BG  = "DCFCE7"
_HEX_BLUE_BG   = "DBEAFE"
_HEX_SECTION_HDR = "2C3E50" # dark slate for stat table header row


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, top="single", bottom="single", left="single", right="single", sz="4", color="E2E8F0") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _para_space(doc: Document, pt: int = 4) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), str(pt * 20))
    pPr.append(spacing)


def _fmt_date(dt_val: datetime | None) -> str:
    if dt_val is None:
        return "—"
    if dt_val.tzinfo is None:
        dt_val = dt_val.replace(tzinfo=timezone.utc)
    return dt_val.strftime("%d %b %Y")


def _section_heading(doc: Document, text: str) -> None:
    """Indigo uppercase section heading with bottom border rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # bottom border
    borders = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), "E0E7FF")
    borders.append(bot)
    pPr.append(borders)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "160")
    spacing.set(qn("w:after"), "80")
    pPr.append(spacing)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _INDIGO
    run.font.all_caps = True


def _sub_heading(doc: Document, text: str) -> None:
    """Bold dark sub-heading for Description, Agenda, etc."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "40")
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _BODY


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Key-value table with shaded label column, matching the HTML kv-table style."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths = (Inches(2.2), Inches(4.0))
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        lc.width = col_widths[0]
        vc.width = col_widths[1]
        lp = lc.paragraphs[0]
        lp.clear()
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = _SLATE
        vp = vc.paragraphs[0]
        vp.clear()
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = _BODY
        _set_cell_bg(lc, _HEX_KV_LABEL)
        bg = _HEX_EVENT_BG if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(vc, bg)


def _stat_table(doc: Document, headers: list[str], values: list[str]) -> None:
    """4-column stat row: dark header row + value row, matching the HTML stat-row."""
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    col_w = Inches(6.2 / len(headers))
    for i, (h, v) in enumerate(zip(headers, values)):
        hc = table.rows[0].cells[i]
        vc = table.rows[1].cells[i]
        hc.width = col_w
        vc.width = col_w
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.clear()
        hr = hp.add_run(h)
        hr.bold = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = _WHITE
        _set_cell_bg(hc, _HEX_SECTION_HDR)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.clear()
        vr = vp.add_run(v)
        vr.font.size = Pt(12)
        vr.bold = True
        vr.font.color.rgb = _BODY
        _set_cell_bg(vc, _HEX_STAT_BOX)


def _summary_stat_table(doc: Document, events_count: int, total_regs: int, total_confirmed: int, total_present: int) -> None:
    """4-box summary grid matching the HTML summary-stats section."""
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    col_w = Inches(1.55)
    stats = [
        ("EVENTS",               str(events_count)),
        ("TOTAL REGISTRATIONS",  str(total_regs)),
        ("CONFIRMED",            str(total_confirmed)),
        ("TOTAL ATTENDANCE",     str(total_present)),
    ]
    for i, (lbl, val) in enumerate(stats):
        lc = table.rows[0].cells[i]
        vc = table.rows[1].cells[i]
        lc.width = col_w
        vc.width = col_w
        # value row (big number)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.clear()
        vr = vp.add_run(val)
        vr.font.size = Pt(20)
        vr.bold = True
        vr.font.color.rgb = _BODY
        _set_cell_bg(vc, _HEX_STAT_BOX)
        # label row
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.clear()
        lr = lp.add_run(lbl)
        lr.font.size = Pt(8)
        lr.bold = True
        lr.font.color.rgb = _MUTED
        _set_cell_bg(lc, _HEX_STAT_BOX)


def _status_badge_color(status: str) -> tuple[str, RGBColor]:
    s = status.upper()
    if s == "COMPLETED":
        return _HEX_GREEN_BG, _GREEN_FG
    if s in ("PUBLISHED", "ACTIVE"):
        return _HEX_BLUE_BG, _BLUE_FG
    return "F3F4F6", _MUTED


def generate_club_report(
    club_name: str,
    department: str | None,
    start_date: date,
    end_date: date,
    events: list[dict[str, Any]],
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Report Header (dark navy block) ──────────────────────────────────────
    # Club name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), _HEX_NAVY)
    pPr.append(shd)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "200"); spacing.set(qn("w:after"), "0")
    pPr.append(spacing)
    r = p.add_run(club_name)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = _WHITE

    for text, size, color in [
        ("Annual Club Activity Report", 13, RGBColor(0xC7, 0xD2, 0xFE)),
        (f"Period: {start_date.strftime('%d %B %Y')} — {end_date.strftime('%d %B %Y')}", 11, RGBColor(0xA5, 0xB4, 0xFC)),
    ]:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2Pr = p2._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear"); shd2.set(qn("w:color"), "auto"); shd2.set(qn("w:fill"), _HEX_NAVY)
        p2Pr.append(shd2)
        sp2 = OxmlElement("w:spacing")
        sp2.set(qn("w:before"), "0"); sp2.set(qn("w:after"), "0")
        p2Pr.append(sp2)
        r2 = p2.add_run(text)
        r2.font.size = Pt(size); r2.font.color.rgb = color

    if department:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3Pr = p3._p.get_or_add_pPr()
        shd3 = OxmlElement("w:shd")
        shd3.set(qn("w:val"), "clear"); shd3.set(qn("w:color"), "auto"); shd3.set(qn("w:fill"), _HEX_NAVY)
        p3Pr.append(shd3)
        sp3 = OxmlElement("w:spacing")
        sp3.set(qn("w:before"), "0"); sp3.set(qn("w:after"), "200")
        p3Pr.append(sp3)
        r3 = p3.add_run(f"Department: {department}")
        r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    _para_space(doc, 8)

    # ── Summary ──────────────────────────────────────────────────────────────
    _section_heading(doc, "Summary")
    _para_space(doc, 4)

    total_regs      = sum(e.get("total_registrations", 0) for e in events)
    total_confirmed = sum(e.get("confirmed_registrations", 0) for e in events)
    total_present   = sum(e.get("attendance_present", 0) for e in events)
    _summary_stat_table(doc, len(events), total_regs, total_confirmed, total_present)
    _para_space(doc, 8)

    if not events:
        p = doc.add_paragraph("No events found for the selected period.")
        p.runs[0].font.color.rgb = _MUTED
        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

    # ── Event Details ─────────────────────────────────────────────────────────
    _section_heading(doc, "Event Details")

    for idx, ev in enumerate(events, 1):
        _para_space(doc, 6)

        # Event title line: "1. EventName   [STATUS]"
        title_p = doc.add_paragraph()
        pPr = title_p._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "80"); spacing.set(qn("w:after"), "40")
        pPr.append(spacing)
        num_run = title_p.add_run(f"{idx}.  ")
        num_run.bold = True; num_run.font.size = Pt(12); num_run.font.color.rgb = _INDIGO
        title_run = title_p.add_run(ev["title"])
        title_run.bold = True; title_run.font.size = Pt(13); title_run.font.color.rgb = _BODY
        status = ev.get("status", "")
        if status:
            _, fg = _status_badge_color(status)
            title_p.add_run("   ")
            badge_run = title_p.add_run(f"  {status}  ")
            badge_run.bold = True; badge_run.font.size = Pt(9); badge_run.font.color.rgb = fg

        # Meta kv table
        meta_rows: list[tuple[str, str]] = [
            ("Date",     f"{_fmt_date(ev.get('start_datetime'))} – {_fmt_date(ev.get('end_datetime'))}"),
            ("Venue",    ev.get("venue") or "—"),
            ("Category", ev.get("category") or "—"),
            ("Type",     ev.get("event_type", "—")),
        ]
        if ev.get("max_participants"):
            meta_rows.append(("Max Participants", str(ev["max_participants"])))
        _kv_table(doc, meta_rows)
        _para_space(doc, 4)

        if ev.get("description"):
            _sub_heading(doc, "Description")
            p = doc.add_paragraph(ev["description"])
            p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = _SLATE
            _para_space(doc, 2)

        if ev.get("agenda"):
            _sub_heading(doc, "Agenda")
            p = doc.add_paragraph(ev["agenda"])
            p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = _SLATE
            _para_space(doc, 2)

        # Registration & Attendance stat boxes
        _sub_heading(doc, "Registration & Attendance")
        regs      = ev.get("total_registrations", 0)
        confirmed = ev.get("confirmed_registrations", 0)
        present   = ev.get("attendance_present", 0)
        att_rate  = f"{present / confirmed * 100:.1f}%" if confirmed else "—"
        _stat_table(
            doc,
            ["Registered", "Confirmed", "Present", "Attendance Rate"],
            [str(regs), str(confirmed), str(present), att_rate],
        )
        _para_space(doc, 4)

        if ev.get("is_team_event"):
            _sub_heading(doc, "Team Information")
            _kv_table(doc, [
                ("Team Event",    "Yes"),
                ("Min Team Size", str(ev.get("team_min_size", "—"))),
                ("Max Team Size", str(ev.get("team_max_size", "—"))),
                ("Total Teams",   str(ev.get("total_teams", 0))),
                ("Avg Team Size", str(ev.get("avg_team_size", 0))),
            ])
            _para_space(doc, 4)

        budget = ev.get("budget", 0) or 0
        spent  = ev.get("spent", 0) or 0
        if budget or spent:
            _sub_heading(doc, "Finance")
            remaining   = budget - spent
            utilization = f"{spent / budget * 100:.1f}%" if budget else "—"
            fmt = lambda n: f"₹{n:,.0f}"
            _kv_table(doc, [
                ("Total Budget",  fmt(budget)),
                ("Amount Spent",  fmt(spent)),
                ("Remaining",     fmt(remaining)),
                ("Utilization",   utilization),
            ])
            _para_space(doc, 4)

        nps = ev.get("nps")
        if nps is not None:
            _sub_heading(doc, "Feedback")
            _kv_table(doc, [("Net Promoter Score (NPS)", str(nps))])
            _para_space(doc, 4)

        if idx < len(events):
            doc.add_page_break()

    # ── Footer ────────────────────────────────────────────────────────────────
    _para_space(doc, 8)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(
        f"Generated on {datetime.now().strftime('%d %B %Y at %I:%M %p')} by ClubHub"
        "  ·  PSG College of Technology Students' Union"
    )
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = _MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
